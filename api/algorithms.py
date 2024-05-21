from fastapi import APIRouter, Body, Depends, HTTPException,UploadFile,File
from fastapi.responses import JSONResponse, FileResponse
from starlette import status
from typing import List, Union, Annotated
from sqlalchemy.orm import Session
from database.db import engine_s
from models.users import *
from models.cases import *
from models.methods import *
import hashlib
import shutil
import os
from api.router import get_session
import docx
from docx.shared import Pt, RGBColor

alg_router = APIRouter(prefix='/algorithm')
@alg_router.get("/api/procedural_position",tags=['Procedural Position'])
def get_procedural_position(DB: Session = Depends(get_session)):
    procedural_pos = DB.query(ProceduralPosition).all()
    if procedural_pos is None:
        return JSONResponse(status_code= 404, content={"message": "Процессуальные положения не найдены"})
    return procedural_pos
@alg_router.get("/individual_involved_info/{id}",tags=['Person Involved'])
def get_individual_involved_info(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == id).first()
    if cases is None:
        return JSONResponse(status_code=404, content={"message":"Дело не найдено!"})
    individual_id = DB.query(Statement).filter(Statement.statement_id == cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individual_id).first()
    residence = DB.query(Passport).filter(Passport.passports_id==individual.FK_passports_id).first().place_of_residence
    info = {
        "first_name":individual.first_name,
        "middle_name":individual.middle_name,
        "last_name":individual.last_name,
        "residence":residence
    }
    return info
@alg_router.post("/person_involved/",tags=['Person Involved'])
def create_person_involved(item:createPersonInvolved,DB:Session = Depends(get_session)):
    person = DB.query(PersonsInvolved).filter(PersonsInvolved.first_name==item.first_name,PersonsInvolved.middle_name==item.middle_name,PersonsInvolved.last_name==item.last_name,PersonsInvolved.residence==item.residence,PersonsInvolved.fk_procedural_position==item.fk_procedural_position).first()
    if person is not None:
        return person.id_persons_involved
    person = PersonsInvolved(first_name=item.first_name,middle_name=item.middle_name,last_name=item.last_name,residence=item.residence,fk_procedural_position=item.fk_procedural_position,notes = item.notes,petition=item.petition)
    try:
        DB.add(person)
        DB.commit()
        DB.refresh(person)
        return person.id_persons_involved
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.get("/person_involved/{id}",tags=['Person Involved'])
def get_person_involved(id:int,DB:Session = Depends(get_session)):
    person = DB.query(PersonsInvolved).filter(PersonsInvolved.id_persons_involved==id).first()
    if person is None:
        return JSONResponse(status_code=404, content={"message":"Человек не найден!"})
    return person
@alg_router.post("/clarif/",tags=['Clarif'])
def clarif_create(item:createClarif,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    clarif = DB.query(Clarif).filter(Clarif.fk_active_case_id == item.case_id).first()
    if clarif is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    clarif = Clarif(fk_active_case_id=item.case_id,date_clarif=item.date_clarif,phone=item.phone,email=item.email,information=item.information)
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(clarif)
        DB.commit()
        DB.refresh(clarif)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/clarif/{caseId}",tags=['Clarif'])
def get_clarif(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    clarif = DB.query(Clarif).filter(Clarif.fk_active_case_id==caseId).first()
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return clarif
@alg_router.put("/clarif/{caseId}",tags=['Clarif'])
def edit_clarif(caseId:int,item:updateClarif,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    clarif = DB.query(Clarif).filter(Clarif.fk_active_case_id==caseId).first()
    if clarif is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    clarif.date_clarif = item.date_clarif
    clarif.email = item.email
    clarif.phone=item.phone
    clarif.information=item.information
    try:
        DB.commit()
        DB.refresh(clarif)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/track/",tags=['Track'])
def track_create(item:createTrack,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    track = DB.query(Track).filter(Track.fk_active_case_id == item.case_id).first()
    if track is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    track = Track(fk_active_case_id=item.case_id,account=item.account,сorrespondence=item.сorrespondence,history=item.history,track_money=item.track_money,logfile=item.logfile,antivirus=item.antivirus)
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(track)
        DB.commit()
        DB.refresh(track)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/track/{caseId}",tags=['Track'])
def get_clarif(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    track = DB.query(Track).filter(Track.fk_active_case_id==caseId).first()
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return track
@alg_router.put("/track/{caseId}",tags=['Track'])
def edit_clarif(caseId:int,item:updateTrack,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    track = DB.query(Track).filter(Track.fk_active_case_id==caseId).first()
    if track is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    track.account = item.account
    track.сorrespondence = item.сorrespondence
    track.history = item.history
    track.track_money=item.track_money
    track.logfile=item.logfile
    track.antivirus=item.antivirus
    try:
        DB.commit()
        DB.refresh(track)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/person_involved_inspect/",tags=['Inspect'])
def create_inspect_person_involved(item:createInspectPersonInvolved,DB:Session = Depends(get_session)):
    inspectperson = DB.query(InspectPersonsInvolved).filter(InspectPersonsInvolved.id_inspect==item.idInspect,InspectPersonsInvolved.id_persons_involved==item.idPerson).first()
    if inspectperson is not None:
        return JSONResponse(status_code=400, content={"message":"Данные уже имеются"})
    inspectperson = InspectPersonsInvolved(id_persons_involved=item.idPerson,id_inspect=item.idInspect)
    try:
        DB.add(inspectperson)
        DB.commit()
        DB.refresh(inspectperson)
        return JSONResponse(status_code=200, content={"message": "Данные добавлены!"})
    except HTTPException:
        return JSONResponse(status_code=500, content={"message":"Ошибка"})
@alg_router.get("/person_involved_inspect/{id}",tags=['Inspect'])
def get_all_inspects_person(id:int,DB:Session = Depends(get_session)):
    inspect_id = DB.query(Inspect).filter(id == Inspect.fk_active_case_id).first().id_inspect
    inspectperson = DB.query(InspectPersonsInvolved).filter(InspectPersonsInvolved.id_inspect==inspect_id).all()
    if len(inspectperson) == 0:
        return JSONResponse(status_code=404,content={"message":"Данные не найдены"})
    all_person = []
    for i in range(len(inspectperson)):
        person = DB.query(PersonsInvolved).filter(PersonsInvolved.id_persons_involved == inspectperson[i].id_persons_involved).first()
        if person.petition == 1:
            petition = 'Нет'
        else:
            petition = 'Да'
        procedural_position_name = DB.query(ProceduralPosition).filter(person.fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        info = {
            "id_persons_involved":person.id_persons_involved,
            "first_name":person.first_name,
            "middle_name":person.middle_name,
            "last_name":person.last_name,
            "residence":person.residence,
            "fk_procedural_position":procedural_position_name,
            "notes":person.notes,
            "petition":petition
        }
        all_person.append(info)
    return all_person
@alg_router.post("/inspect/",tags=['Inspect'])
def inspect_create(item:createInspect,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id == item.case_id).first()
    if inspect is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    inspect = Inspect(fk_active_case_id=item.case_id,date_inspect=item.date_inspect,begin_inspect=item.begin_inspect,end_inspect=item.end_inspect,message_inspect=item.message_inspect,from_message_inspect=item.from_message_inspect,place_of_inspect=item.place_of_inspect,technical_means=item.technical_means,conditions=item.conditions,establish=item.establish,photography=item.photography,sized_items=item.sized_items,items_for_inspect=item.items_for_inspect,familiarization=item.familiarization,inspect_exam=item.inspect_exam)
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(inspect)
        DB.commit()
        DB.refresh(inspect)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/inspect/{caseId}",tags=['Inspect'])
def get_inspect(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id==caseId).first()
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return inspect
@alg_router.put("/inspect/{caseId}",tags=['Inspect'])
def edit_inspect(caseId:int,item:updateInspect,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id==caseId).first()
    if inspect is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    inspect.date_inspect = item.date_inspect
    inspect.begin_inspect = item.begin_inspect
    inspect.end_inspect=item.end_inspect
    inspect.message_inspect=item.message_inspect
    inspect.from_message_inspect=item.from_message_inspect
    inspect.place_of_inspect=item.place_of_inspect
    inspect.technical_means=item.technical_means
    inspect.conditions=item.conditions
    inspect.establish=item.establish
    inspect.photography=item.photography
    inspect.sized_items=item.sized_items
    inspect.items_for_inspect=item.items_for_inspect
    inspect.familiarization=item.familiarization
    inspect.inspect_exam=item.inspect_exam
    try:
        DB.commit()
        DB.refresh(inspect)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/reclaim/{id}/{caseId}/upload",tags=['Reclaim'])
def reclaim_upload(id:int,caseId:int,passp:UploadFile=File(...),inc:UploadFile=File(...),detail:UploadFile=File(...),DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    reclaim = DB.query(Reclaim).filter(Reclaim.fk_active_case_id==caseId).first()
    if reclaim is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    passp.filename = passp.filename.lower()
    inc.filename = inc.filename.lower()
    detail.filename = detail.filename.lower()
    if os.path.exists(f'media/{id}/cases/{caseId}'):
        path_psr=f'media/{id}/cases/{caseId}/reclaim/passport/{passp.filename}'
        with open(path_psr,'wb+') as buffer:
            shutil.copyfileobj(passp.file,buffer)
        path_inc=f'media/{id}/cases/{caseId}/reclaim/income/{inc.filename}'
        with open(path_inc,'wb+') as buffer:
            shutil.copyfileobj(inc.file,buffer)
        path_det=f'media/{id}/cases/{caseId}/reclaim/detailing/{detail.filename}'
        with open(path_det,'wb+') as buffer:
            shutil.copyfileobj(detail.file,buffer)
        reclaim = Reclaim(fk_active_case_id=caseId,passport=path_psr,income=path_inc,detailing=path_det)
        try:
            DB.add(reclaim)
            DB.commit()
            DB.refresh(reclaim)
            return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
        except Exception as e:
            raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
    else:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
@alg_router.get("/get_reclaim/{caseId}",tags=['Reclaim'])
def get_reclaim(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    reclaim = DB.query(Reclaim).filter(Reclaim.fk_active_case_id==caseId).first()
    if reclaim is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return reclaim
@alg_router.get("/get_reclaim/{caseId}/{anyDocs}",tags=['Reclaim'])
def get_reclaim_passport(caseId:int,anyDocs:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    reclaim = DB.query(Reclaim).filter(Reclaim.fk_active_case_id==caseId).first()
    if reclaim is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    if anyDocs == 'passport':
        return FileResponse(reclaim.passport)
    elif anyDocs == 'income':
        return FileResponse(reclaim.income)
    elif anyDocs == 'detailing':
        return FileResponse(reclaim.detailing)
    else:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    
@alg_router.post("/extract/{id}/{caseId}/upload",tags=['Extract'])#
def extract_upload(id:int,caseId:int,mon_ord:UploadFile=File(...),scrsht:UploadFile=File(...),detail:UploadFile=File(...),DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    extraxt = DB.query(Extract).filter(Extract.fk_active_case_id==caseId).first()
    if extraxt is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    mon_ord.filename = mon_ord.filename.lower()
    scrsht.filename = scrsht.filename.lower()
    detail.filename = detail.filename.lower()
    if os.path.exists(f'media/{id}/cases/{caseId}'):
        path_mon_ord=f'media/{id}/cases/{caseId}/extract/money_orders/{mon_ord.filename}'
        with open(path_mon_ord,'wb+') as buffer:
            shutil.copyfileobj(mon_ord.file,buffer)
        path_scrsht=f'media/{id}/cases/{caseId}/extract/screenshot/{scrsht.filename}'
        with open(path_scrsht,'wb+') as buffer:
            shutil.copyfileobj(scrsht.file,buffer)
        path_det=f'media/{id}/cases/{caseId}/extract/detailing/{detail.filename}'
        with open(path_det,'wb+') as buffer:
            shutil.copyfileobj(detail.file,buffer)
        extraxt = Extract(fk_active_case_id=caseId,money_orders=path_mon_ord,screenshot=path_scrsht,detailing=path_det)
        try:
            DB.add(extraxt)
            DB.commit()
            DB.refresh(extraxt)
            return JSONResponse(status_code= 200, content={"message": "Файлы успешно загружены"})
        except Exception as e:
            raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
    else:
        return JSONResponse(status_code= 404, content={"message": "Добавление файлов невозможно!"})
@alg_router.get("/get_extract/{caseId}",tags=['Extract'])
def get_extract(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    extract = DB.query(Extract).filter(Extract.fk_active_case_id==caseId).first()
    if extract is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return extract
@alg_router.get("/get_extract/{caseId}/{anyDocs}",tags=['Extract'])
def get_extract_docs(caseId:int,anyDocs:str,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    extract = DB.query(Extract).filter(Extract.fk_active_case_id==caseId).first()
    if extract is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    if anyDocs == 'screenshot':
        return FileResponse(extract.screenshot)
    elif anyDocs == 'money_orders':
        return FileResponse(extract.money_orders)
    elif anyDocs == 'detailing':
        return FileResponse(extract.detailing)
    else:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})

@alg_router.post("/check/",tags=['Check'])
def check_create(item:createCheck,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    check = DB.query(Check).filter(Check.fk_active_case_id == item.case_id).first()
    if check is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    check = Check(fk_active_case_id=item.case_id,checks=item.checks)
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(check)
        DB.commit()
        DB.refresh(check)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/check/{caseId}",tags=['Check'])
def get_check(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    check = DB.query(Check).filter(Check.fk_active_case_id==caseId).first()
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return check
@alg_router.put("/check/{caseId}",tags=['Check'])
def edit_check(caseId:int,item:updateCheck,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    check = DB.query(Check).filter(Check.fk_active_case_id==caseId).first()
    if check is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    check.checks = item.checks
    try:
        DB.commit()
        DB.refresh(check)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/expertise/",tags=['Expertise'])
def expertise_create(item:createExpertise,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id == item.case_id).first()
    if expertise is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    expertise = Expertise(fk_active_case_id=item.case_id,app_comp=item.app_comp,prog_comp=item.prog_comp,inf_comp=item.inf_comp,comp_net=item.comp_net)
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(expertise)
        DB.commit()
        DB.refresh(expertise)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/expertise/{caseId}",tags=['Expertise'])
def get_expertise(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id==caseId).first()
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return expertise
@alg_router.put("/expertise/{caseId}",tags=['Expertise'])
def edit_expertise(caseId:int,item:updateExpertise,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    expertise = DB.query(Expertise).filter(Expertise.fk_active_case_id==caseId).first()
    if expertise is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    expertise.app_comp = item.app_comp
    expertise.prog_comp = item.prog_comp
    expertise.inf_comp = item.inf_comp
    expertise.comp_net = item.comp_net
    try:
        DB.commit()
        DB.refresh(expertise)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
    
@alg_router.post("/quest/",tags=['Quest'])
def quest_create(item:createQuest,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    quest = DB.query(Quest).filter(Quest.fk_active_case_id == item.case_id).first()
    if quest is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    quest = Quest(fk_active_case_id=item.case_id,have=item.have,reg=item.reg,know=item.know,have_bank=item.have_bank,connect_online=item.connect_online,know_cards=item.know_cards,pay_site=item.pay_site)
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(quest)
        DB.commit()
        DB.refresh(quest)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/quest/{caseId}",tags=['Quest'])
def get_quest(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    quest = DB.query(Quest).filter(Quest.fk_active_case_id==caseId).first()
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return quest
@alg_router.put("/quest/{caseId}",tags=['Quest'])
def edit_quest(caseId:int,item:updateQuest,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    quest = DB.query(Quest).filter(Quest.fk_active_case_id==caseId).first()
    if quest is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    quest.have = item.have
    quest.reg = item.reg
    quest.know = item.know
    quest.have_bank = item.have_bank
    quest.connect_online = item.connect_online
    quest.know_cards = item.know_cards
    quest.pay_site = item.pay_site
    try:
        DB.commit()
        DB.refresh(quest)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})

@alg_router.post("/set/",tags=['Set'])
def set_create(item:createSet,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    set = DB.query(Set).filter(Set.fk_active_case_id == item.case_id).first()
    if set is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    set = Set(fk_active_case_id=item.case_id,criminal=item.criminal)
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(set)
        DB.commit()
        DB.refresh(set)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/set/{caseId}",tags=['Set'])
def get_set(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    set = DB.query(Set).filter(Set.fk_active_case_id==caseId).first()
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return set
@alg_router.put("/set/{caseId}",tags=['Set'])
def edit_set(caseId:int,item:updateSet,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    set = DB.query(Set).filter(Set.fk_active_case_id==caseId).first()
    if set is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    set.criminal=item.criminal
    try:
        DB.commit()
        DB.refresh(set)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})   
@alg_router.post("/witnes/",tags=['Witnes'])
def witnes_create(item:createWitnes,DB:Session = Depends(get_session)):
    user = DB.query(User).filter(User.id == item.user_id).first()
    if user is None:
        return JSONResponse(status_code= 404, content={"message": "Пользователь не найден"})
    cases = DB.query(Case).filter(Case.active_cases_id == item.case_id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    if cases.fk_user_id != user.id:
        return JSONResponse(status_code= 400, content={"message": "Этот пользователь не может редактировать данное дело"})
    witnes = DB.query(Witnes).filter(Witnes.fk_active_case_id == item.case_id).first()
    if witnes is not None:
        return JSONResponse(status_code= 400, content={"message": "Данные уже добавлены"})
    witnes = Witnes(fk_active_case_id=item.case_id,witnes=item.witnes)
    if witnes is None:
        return JSONResponse(status_code= 404, content={"message": "Произошла ошибка!"})
    try:
        DB.add(witnes)
        DB.commit()
        DB.refresh(witnes)
        return JSONResponse(status_code= 200, content={"message": "Данные добавлены!"})
    except Exception as e:
        raise HTTPException(status_code=500, detail= f"Произошла ошибка при добавлении объекта")
@alg_router.get("/witnes/{caseId}",tags=['Witnes'])
def get_witnes(caseId:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    witnes = DB.query(Witnes).filter(Witnes.fk_active_case_id==caseId).first()
    if witnes is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    return witnes
@alg_router.put("/witnes/{caseId}",tags=['Witnes'])
def edit_witnes(caseId:int,item:updateWitnes,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id == caseId).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    witnes = DB.query(Witnes).filter(Witnes.fk_active_case_id==caseId).first()
    if witnes is None:
        return JSONResponse(status_code= 404, content={"message": "Данные не найдены"})
    witnes.witnes=item.witnes
    try:
        DB.commit()
        DB.refresh(witnes)
        return JSONResponse(status_code=200, content={"message": "Данные обновлены"})
    except HTTPException:
        return JSONResponse(status_code=404, content={"message":"Ошибка"})
@alg_router.post("/request/mobile_operator/download/{id}",tags=['Request'])
def create_mobile_request_document(id:int,item:requestMobile,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    individuals_id=DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individuals_id).first()
    individualFIO = individual.first_name + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    user = DB.query(User).filter(User.id==cases.fk_user_id).first()
    userFIO = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    individual_sex = DB.query(Passport).filter(Passport.passports_id==individual.FK_passports_id).first().sex
    department = DB.query(Department).filter(user.fk_department_id == Department.departments_id).first()
    name_department = department.name_department
    victim = 'потерпевшей'
    if individual_sex=='Муж':
        victim = 'потерпевшему'
    date_case = item.date_case.strftime("%d.%m.%Y")
    date_incidient = item.date_incidient.strftime("%d.%m.%Y")
    src = 'documents/requestmobile.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    doc = docx.Document(dst+'requestmobile.docx')
    dictionary = {"СОВЕТСКИЙ":name_department.upper(),"Советский":department.name_department,"Челябинск":department.city,"ЧЕЛЯБИНСК":department.city.upper(),"ул. Монакова 2":department.street+', '+department.house_number,"454091":department.postal_code,"1220175009800000":cases.number_cases,"30.01.2022":date_case,"Сидоровой А.В.":individualFIO,"13 часов":item.time[0]+item.time[1]+' часов',"31 минуту":item.time[3]+item.time[4]+' минуту',"21 января 2022":date_incidient,"8-499-769-34-13":item.phone,"11 минут 24 секунды":item.duration,"ivanova@mvd.ru":user.email,"И.И. Иванова":userFIO,"потерпевшей":victim,"Старший следователь":user.appointment,"старший лейтенант":user.users_rank}
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
    doc.save(dst+'requestmobile.docx')
    return FileResponse(dst+'requestmobile.docx',filename="Запрос мобильному оператору.docx",media_type='application/octet-stream')
@alg_router.post("/request/bank/download/{id}",tags=['Request'])
def create_bank_request_document(id:int,item:requestBank,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code= 404, content={"message": "Дело не найдено"})
    individuals_id=DB.query(Statement).filter(Statement.statement_id==cases.FK_statement_id).first().fk_individuals_id
    individual = DB.query(Invidivual).filter(Invidivual.individuals_id==individuals_id).first()
    individualFIO = individual.first_name + ' ' + individual.middle_name[0]+'.' + individual.last_name[0]+'.'
    user = DB.query(User).filter(User.id==cases.fk_user_id).first()
    userFIO = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    mainUser = DB.query(User).filter(User.FK_user_roles_id==4,User.fk_department_id==user.fk_department_id).first()
    mainUserFIO = mainUser.middle_name[0]+'.' + mainUser.last_name[0]+'.'+ ' ' + mainUser.first_name
    department = DB.query(Department).filter(user.fk_department_id == Department.departments_id).first()
    name_department = department.name_department
    src = 'documents/requestbank.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    doc = docx.Document(dst+'requestbank.docx')
    dictionary = {"СОВЕТСКИЙ":name_department.upper(), "Советский":department.name_department,
                  "Челябинск":department.city,"ЧЕЛЯБИНСК":department.city.upper(),
                  "ул. Монакова 2":department.street+', '+department.house_number,"454091":department.postal_code,
                  "1220175009800000":cases.number_cases,
                  "АО «Альфа Банк»":item.bank,
                  "Сидоровой А.В.":individualFIO,"ivanovamvd@.ru":user.email,"И.И. Иванова":userFIO,"Старший следователь":user.appointment,"старший лейтенант":user.users_rank,
                  "К.А.Борисов":mainUserFIO, "подполковник":mainUser.users_rank,"И.И.  Ивановой":userFIO}
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
    doc.save(dst+'requestbank.docx')
    return FileResponse(dst+'requestbank.docx',filename="Запрос в банк.docx",media_type='application/octet-stream')
@alg_router.get("/documents/inspect/{id}",tags=['Inspect'])
def get_inspect_document(id:int,DB:Session = Depends(get_session)):
    cases = DB.query(Case).filter(Case.active_cases_id==id).first()
    if cases is None:
        return JSONResponse(status_code=404,content={"message":"Дело не найдено"})
    department_id = DB.query(User).filter(User.id==cases.fk_user_id).first().fk_department_id
    department = DB.query(Department).filter(Department.departments_id==department_id).first()
    inspect = DB.query(Inspect).filter(Inspect.fk_active_case_id == cases.active_cases_id).first()
    date_inspect = inspect.date_inspect.strftime("%d.%m.%Y")
    user = DB.query(User).filter(User.id == cases.fk_user_id).first()
    userFIO_1 = user.first_name + ' ' + user.middle_name[0]+'.' + user.last_name[0]+'.'
    userFIO_2 = user.middle_name[0]+'.' + user.last_name[0]+'.'+ ' ' + user.first_name
    inspect_people = DB.query(InspectPersonsInvolved).filter(inspect.id_inspect == InspectPersonsInvolved.id_inspect).all()
    p1 = ''
    p2 = ''
    pet = ''
    notes = ''
    otherPerson = ''
    for i in range(len(inspect_people)):
        person = DB.query(PersonsInvolved).filter(PersonsInvolved.id_persons_involved == inspect_people[i].id_persons_involved).first()
        procedural_position_name = DB.query(ProceduralPosition).filter(person.fk_procedural_position==ProceduralPosition.id_procedural_position).first().name_position
        if procedural_position_name == 'Понятой':
            if len(p1) == 0:
                p1 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ person.residence
            else:
                p2 = person.first_name+' '+person.middle_name+' '+person.last_name+'. Место жительства: '+ person.residence
        elif procedural_position_name == 'Эксперт':
            expFIO = person.first_name+' '+person.middle_name+' '+person.last_name
            exp = 'Эксперт'
            state = '57'
        elif procedural_position_name == 'Специалист':
            expFIO = person.first_name+' '+person.middle_name+' '+person.last_name
            exp = 'Специалист'
            state = '58'
        else:
            otherPerson += procedural_position_name + ' '+ person.first_name + ' ' + person.middle_name + ' ' + person.last_name
        if person.petition == 2:
            pet += procedural_position_name+' '+person.first_name + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.'
        if len(person.notes) != 0:
            notes += procedural_position_name+' '+person.first_name + ' '+person.middle_name[0]+'.'+person.last_name[0]+'.' +' Замечание - ' +person.notes+'.'
    if pet == '':
        pet = 'Следователя'+userFIO_1
    if len(notes) != 0:
        notes ='сделали - ' + notes
    else:
        notes = 'не сделали.'
    src = 'documents/omp.docx'
    dst = f'media/{cases.fk_user_id}/cases/{cases.active_cases_id}/documents/'
    shutil.copy(src, dst, follow_symlinks=True)
    doc = docx.Document(dst+'omp.docx')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    dictionary = {
        "othp": otherPerson,
        "inex": inspect.inspect_exam.lower(),
        "pet": pet,
        "notes": notes,
        "p1":p1,
        "p2":p2,
        "Exs":exp,
        "exFIO":expFIO,
        "st": state,
        "App": user.appointment,
        "ur": user.users_rank,
        "FIO1": userFIO_1,
        "FIO2": userFIO_2,
        "(мс)": "г. "+department.city,
        "Домп": date_inspect,
        "омпнч": inspect.begin_inspect[0]+inspect.begin_inspect[1],
        "онм": inspect.begin_inspect[3]+inspect.begin_inspect[4],
        "ооч": inspect.end_inspect[0]+inspect.end_inspect[1],
        "оом": inspect.end_inspect[3]+inspect.end_inspect[4],
        "окс": inspect.from_message_inspect.lower(),
        "очс": inspect.message_inspect.lower(),
        "кпс": inspect.place_of_inspect.lower(),
        "tm": inspect.technical_means.lower(),
        "уос": inspect.conditions.lower(),
        "оуе": inspect.establish.lower(),
        "фва": inspect.photography.lower(),
        "si": inspect.sized_items.lower(),
        "ifi": inspect.items_for_inspect.lower(),
        "fam":inspect.familiarization.lower()
    }
    all_tables = doc.tables
    
    for i in dictionary:
        for p in doc.paragraphs:
            if p.text.find(i) >= 0:
                p.text = p.text.replace(i,dictionary[i])
        for k, table in enumerate(all_tables):
            for j, row in enumerate(table.rows):
                for cell in row.cells:
                    if cell.text.find(i) >= 0:
                        cell.text = cell.text.replace(i,dictionary[i])
    doc.save(dst+'omp.docx')
    return FileResponse(dst+'omp.docx',filename="Протокол осмотра места происшествия.docx",media_type='application/octet-stream')